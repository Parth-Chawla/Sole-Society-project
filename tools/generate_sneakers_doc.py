from pathlib import Path
import re
from docx import Document

ROOT = Path(r"c:\Users\Mohakk\OneDrive\Desktop\Sole Society")
SEED = ROOT / "db" / "seeds" / "001_seed_sneakers.sql"
OUT_DIR = ROOT / "docs"
OUT_DIR.mkdir(exist_ok=True)
OUT_FILE = OUT_DIR / "sneakers_table.docx"

text = SEED.read_text(encoding="utf-8")

# find first INSERT INTO sneakers (...) VALUES (...);
m = re.search(r"INSERT\s+INTO\s+sneakers\s*\(([^)]+)\)\s*VALUES\s*(.+?);", text, re.IGNORECASE | re.DOTALL)
if not m:
    raise SystemExit("No INSERT INTO sneakers(...) found in seed file.")

cols = [c.strip().strip('`"') for c in m.group(1).split(",")]
# remove 'stock' column if present
cols_no_stock = [c for c in cols if c.lower() != "stock"]

values_part = m.group(2)
tuples = re.findall(r"\((.*?)\)(?:\s*,\s*|$)", values_part, re.DOTALL)

rows = []
for t in tuples:
    parts = []
    cur = ""
    in_q = False
    i = 0
    while i < len(t):
        ch = t[i]
        if ch == "'" and (i == 0 or t[i-1] != "\\"):
            in_q = not in_q
            cur += ch
        elif ch == "," and not in_q:
            parts.append(cur.strip())
            cur = ""
        else:
            cur += ch
        i += 1
    if cur:
        parts.append(cur.strip())
    parts = [p.strip().strip("'").strip() for p in parts]

    # remove stock value if original cols included it
    if "stock" in [c.lower() for c in cols]:
        idx_stock = [c.lower() for c in cols].index("stock")
        parts = [v for idx, v in enumerate(parts) if idx != idx_stock]

    # ensure row has exactly same number of columns as cols_no_stock
    target_len = len(cols_no_stock)
    if len(parts) > target_len:
        parts = parts[:target_len]
    elif len(parts) < target_len:
        parts = parts + [""] * (target_len - len(parts))

    rows.append(parts)

# Build docx
doc = Document()
doc.add_heading("Sneakers Seed Table (stock removed)", level=2)
table = doc.add_table(rows=1, cols=len(cols_no_stock))
hdr_cells = table.rows[0].cells
for i, h in enumerate(cols_no_stock):
    hdr_cells[i].text = h

for r in rows:
    row_cells = table.add_row().cells
    for i in range(len(cols_no_stock)):
        row_cells[i].text = r[i]

doc.save(OUT_FILE)
print(f"Generated: {OUT_FILE}")